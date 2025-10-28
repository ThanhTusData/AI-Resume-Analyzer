"""
DOCX Parser
Extracts text from Word documents
"""

from docx import Document
from loguru import logger

from src.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """Parser for DOCX resume files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['docx', 'doc']
    
    def parse(self, file_path: str) -> str:
        """Parse DOCX file"""
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine
            all_text = paragraphs + table_texts
            text = "\n".join(all_text)
            
            logger.info(f"Parsed DOCX: {len(text)} chars")
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}")
            raise